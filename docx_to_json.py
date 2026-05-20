#!/usr/bin/env python3
"""
Конвертер рецептов из Word (.docx) в JSON для загрузки на gastronom.ru.

Заточен под формат:
    Имя автора
    **Название рецепта**
    Описание
    [таблица метаданных с порциями, временем и т.д.]
    ИНГРЕДИЕНТЫ
    · Ингредиент 1 — количество единица (пояснение)
    · Ингредиент 2 — количество единица
    РЕЦЕПТ ПРИГОТОВЛЕНИЯ
    Шаг 1
    Текст шага 1
    Шаг 2
    Текст шага 2
    КСТАТИ
    Текст блока "Кстати"
    ХОЗЯЙКЕ НА ЗАМЕТКУ
    Текст блока "Хозяйке на заметку"

Использование:
    python3 docx_to_json.py путь/к/папке/с/рецептами
    python3 docx_to_json.py recipe.docx

Результат: recipes.json + папка images/ с картинками из docx.
"""

import json
import re
import sys
import zipfile
from pathlib import Path
from typing import Optional

try:
    from docx import Document
except ImportError:
    print("Нужно установить библиотеку python-docx. Запусти:")
    print("    pip install python-docx")
    sys.exit(1)


# Ключевые слова для поиска секций (нижний регистр)
SECTION_INGREDIENTS = ["ингредиенты", "состав", "продукты"]
SECTION_STEPS = ["рецепт приготовления", "способ приготовления",
                 "пошаговый рецепт", "пошагово", "шаги приготовления",
                 "приготовление"]
SECTION_TIP = ["кстати"]
SECTION_HOUSEWIFE = ["хозяйке на заметку", "совет"]

# Префиксы строк SEO-блока, которые надо пропускать в шапке документа.
# Если строка начинается с такого префикса — это служебная информация, не описание.
SEO_PREFIXES = [
    "запрос 1", "запрос 2", "запрос 3",
    "структура -", "структура –",
    "ключевые фразы", "ключевые слова",
    "примечание -", "примечание –", "примечание:",
    "примеры сайтов",
    "тип:",  # редакторская пометка "Тип: сезонный" и т.п.
]
# Подстроки, по которым определяется, что текст — часть SEO-блока
SEO_INNER_MARKERS = [
    "ключевые слова -", "ключевые слова –",
    "ключевые фразы -", "ключевые фразы –",
    "@povar.ru", "@аймкук", "@азбука рецептов",
]


def is_group_line(text: str) -> Optional[str]:
    """Возвращает название группы если строка — заголовок группы ингредиентов (кончается на ':').
    Например: 'Для теста:' → 'Для теста'. Иначе None."""
    t = text.strip()
    if not t.endswith(":"):
        return None
    name = t[:-1].strip()
    if not name or len(name) > 60:
        return None
    # Не должно быть похоже на шаг (начинается с цифры и точки)
    if re.match(r"^\d+[\.\)]", name):
        return None
    # Не должно быть похоже на URL
    if name.startswith("http"):
        return None
    return name + ":"


def is_seo_line(text: str) -> bool:
    """Похожа ли строка на SEO-метаданные, которые не нужно публиковать."""
    t = text.lower().strip()
    # URL — почти всегда мусор (антиплагиат, ссылки на источники)
    if t.startswith("http://") or t.startswith("https://"):
        return True
    for p in SEO_PREFIXES:
        if t.startswith(p):
            return True
    for m in SEO_INNER_MARKERS:
        if m in t:
            return True
    return False

# Маркеры строк-шагов. Распознаём:
#   "Шаг 1", "Шаг N"
#   "1.", "2.", "3." — отдельной строкой
#   "1)", "2)", "3)" — отдельной строкой
#   "1. Возьмите чеснок..." — в начале строки с текстом
STEP_HEADER_ONLY_RE = re.compile(r"^\s*шаг\s+\d+\s*$", re.IGNORECASE)
STEP_NUMBER_ONLY_RE = re.compile(r"^\s*\d+\s*[\.\)]\s*$")
# Пробел между точкой/скобкой и текстом — необязательный (бывает "1.Нагрейте")
STEP_INLINE_RE = re.compile(r"^\s*(\d+)\s*[\.\)]\s*(.+)$")

# Маркеры маркированного списка ингредиентов
LIST_MARKERS = "·•●○◦▪▫⁃-—*"


def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s.strip().lower()).rstrip(":")


def is_section(text: str, keywords: list) -> bool:
    n = normalize(text)
    return any(n == k or n.startswith(k) for k in keywords)


def collect_paragraphs(doc) -> list:
    """
    Собираем все непустые абзацы и таблицы документа В ПОРЯДКЕ их появления.
    """
    items = []

    # Используем высокоуровневое API: iter_inner_content() даёт параграфы и таблицы по порядку
    # (доступно в python-docx 1.0+). Если его нет — собираем по document.element.body.
    try:
        body_iter = doc.iter_inner_content()
    except AttributeError:
        # Fallback: руками сопоставляем XML с параграфами/таблицами
        body_iter = []
        para_idx = 0
        tbl_idx = 0
        for child in doc.element.body.iterchildren():
            tag = child.tag.split("}")[-1]
            if tag == "p":
                if para_idx < len(doc.paragraphs):
                    body_iter.append(doc.paragraphs[para_idx])
                    para_idx += 1
            elif tag == "tbl":
                if tbl_idx < len(doc.tables):
                    body_iter.append(doc.tables[tbl_idx])
                    tbl_idx += 1

    from docx.text.paragraph import Paragraph
    from docx.table import Table

    for block in body_iter:
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            # Shift+Enter внутри Word создаёт мягкий перенос (\n) в рамках одного
            # параграфа. Разбиваем на отдельные строки — иначе вся группа ингредиентов
            # ("Для теста:\nМука – 200 г\n...") попадёт в одно поле name.
            for sub in text.split('\n'):
                sub = sub.strip()
                if sub:
                    items.append(("p", sub))
        elif isinstance(block, Table):
            rows = []
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                items.append(("table", rows))

    return items


def parse_table_meta(rows: list) -> dict:
    """
    Из таблицы метаданных вытаскиваем порции, время, доп. время и прочие поля.
    Таблица имеет вид: [["ПОРЦИИ", "6"], ["ВРЕМЯ", "30 минут"], ...]
    """
    meta = {}
    for row in rows:
        if len(row) < 2:
            continue
        key = normalize(row[0]).replace("*", "").strip()
        value = row[1].strip()
        if not key or not value:
            continue

        if "порц" in key:
            m = re.search(r"\d+", value)
            if m:
                meta["servings"] = int(m.group())
        elif "дополнительное время" in key or "доп. время" in key or "доп время" in key:
            t = extract_minutes(value)
            if t is not None:
                meta["extra_time"] = t
        elif key == "время" or "общее время" in key:
            t = extract_minutes(value)
            if t is not None:
                meta["cooking_time"] = t
        elif "сложность" in key:
            meta["difficulty"] = value
        elif "повод" in key:
            meta["occasion"] = value
        elif "категория" in key:
            meta["category"] = value
        elif "устройство" in key:
            meta["device"] = value
        elif "технология" in key:
            meta["technology"] = value
        elif "кухня" in key:
            meta["cuisine"] = value
        elif "тип вегетарианства" in key:
            meta["vegetarian_type"] = value
        elif "тип питания" in key:
            meta["diet_type"] = value
        elif "тип блюда" in key:
            meta["dish_type"] = value
        elif "основной ингредиент" in key:
            meta["main_ingredient"] = value
    return meta


def extract_minutes(text: str) -> Optional[int]:
    """Преобразуем '30 мин', '2 ч', '2ч', '2 ч.', '1 час', '1,5 часа', '1 час 20 мин' → минуты."""
    total = 0
    # Часы: полное слово "час..." или сокращение "ч" / "ч." (не часть другого слова)
    h = re.search(r"(\d+(?:[.,]\d+)?)\s*(?:час|ч\.?(?=\s|$|,|;|\d))", text, re.IGNORECASE)
    if h:
        total += int(float(h.group(1).replace(",", ".")) * 60)
    # Минуты: "мин..." (минут, минуты, мин., мин)
    m = re.search(r"(\d+)\s*мин", text, re.IGNORECASE)
    if m:
        total += int(m.group(1))
    return total if total > 0 else None


def parse_ingredient_line(line: str) -> Optional[dict]:
    """
    Разбираем строку вида:
        "Сазан — 2 кг (одна тушка)"
        "Лимон — 1 шт."
        "Соль крупная — 1 ст. л."
        "Соль по вкусу"
        "Творог 500 г"
        "Майонез — ½ ст. л."
    """
    # Убираем маркеры списка
    line = re.sub(rf"^\s*[{re.escape(LIST_MARKERS)}]\s*", "", line).strip()
    if not line:
        return None

    # Заменяем unicode-дроби на десятичные числа, чтобы регулярка их понимала
    UNICODE_FRACTIONS = {
        "½": "0.5", "⅓": "0.33", "⅔": "0.67",
        "¼": "0.25", "¾": "0.75",
        "⅕": "0.2",  "⅖": "0.4",  "⅗": "0.6",  "⅘": "0.8",
        "⅙": "0.17", "⅚": "0.83",
        "⅛": "0.125","⅜": "0.375","⅝": "0.625","⅞": "0.875",
    }
    for char, replacement in UNICODE_FRACTIONS.items():
        # Если дробь идёт сразу после цифры — это смешанное число (1½ → 1.5)
        line = re.sub(r"(\d+)\s*" + re.escape(char),
                      lambda m: str(int(m.group(1)) + float(replacement)), line)
        # Иначе просто заменяем
        line = line.replace(char, replacement)

    # Отделяем пояснение в скобках, чтобы не мешало парсингу
    note_match = re.search(r"\(([^)]+)\)", line)
    note = note_match.group(1).strip() if note_match else ""
    line_clean = re.sub(r"\([^)]*\)", "", line).strip()
    # Висячая запятая после скобок
    line_clean = re.sub(r"[,\s]+$", "", line_clean)

    # Список единиц (для регулярки) — упорядочен по убыванию длины
    units_canonical = [
        "ст.л.", "ст. л.", "ст л", "ст.л",
        "ч.л.", "ч. л.", "ч л", "ч.л",
        "на кончике ножа", "по вкусу",
        "кг", "мл", "л", "г",
        "шт.", "шт",
        "стакан", "стакана", "стаканов",
        "пучок", "пучка",
        "ломтик", "ломтика",
        "головка", "головки",
        "веточка", "веточки",
        "зубчик", "зубчика", "зубчиков",
        "щепотка", "щепотки",
    ]
    units_pattern = "|".join(re.escape(u) for u in units_canonical)

    # Число (целое, дробное через , или ., обыкновенная дробь, диапазон)
    num_pattern = r"\d+(?:[.,/]\d+)?(?:\s*[-–]\s*\d+(?:[.,/]\d+)?)?"

    def clean_name(name: str, amount: str) -> str:
        """Убирает хвостовую цифру из имени, если она — нижняя граница диапазона.
        Пример: "Чеснок 3" + amount="4" → "Чеснок" (из "Чеснок 3 — 4 зубчика")."""
        if amount:
            name = re.sub(r"\s+\d+[\d.,]*\s*$", "", name).strip()
        return name

    # Шаблон 1: "название — число единица"
    # Разделитель — только длинное тире (— или –), НЕ короткий дефис.
    # Иначе строка типа "Чеснок 3-4 зубчика" интерпретируется как "Чеснок 3 — 4 зубчика".
    pattern1 = rf"^(.+?)\s*[—–]\s*({num_pattern})\s*({units_pattern})?\s*$"
    m = re.match(pattern1, line_clean, re.IGNORECASE)
    if m:
        name = clean_name(m.group(1).strip(), m.group(2).strip())
        return {
            "name": name,
            "amount": m.group(2).strip(),
            "unit": (m.group(3) or "").strip(),
            "note": note,
        }

    # Шаблон 2: "название по вкусу" / "название на кончике ножа"
    pattern2 = rf"^(.+?)\s+(по вкусу|на кончике ножа)\s*$"
    m = re.match(pattern2, line_clean, re.IGNORECASE)
    if m:
        return {
            "name": m.group(1).strip(),
            "amount": "",
            "unit": m.group(2).strip(),
            "note": note,
        }

    # Шаблон 3: "название число единица" (без тире)
    pattern3 = rf"^(.+?)\s+({num_pattern})\s*({units_pattern})?\s*$"
    m = re.match(pattern3, line_clean, re.IGNORECASE)
    if m:
        return {
            "name": m.group(1).strip(),
            "amount": m.group(2).strip(),
            "unit": (m.group(3) or "").strip(),
            "note": note,
        }

    # Не разобрали — кладём всё в name
    return {"name": line_clean, "amount": "", "unit": "", "note": note}


def extract_images_from_docx(docx_path: Path, output_dir: Path) -> list:
    """Достаём картинки из .docx (zip-архив)."""
    images = []
    output_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path) as z:
        media = sorted(n for n in z.namelist() if n.startswith("word/media/"))
        for i, name in enumerate(media, 1):
            ext = Path(name).suffix.lower()
            if ext not in {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp"}:
                continue
            out = output_dir / f"image_{i}{ext}"
            with z.open(name) as src, open(out, "wb") as dst:
                dst.write(src.read())
            images.append(out)
    return images


def convert_docx(docx_path: Path, images_root: Path, external_images: list = None) -> Optional[dict]:
    print(f"\n→ Обрабатываю: {docx_path.name}")
    doc = Document(docx_path)
    items = collect_paragraphs(doc)
    if not items:
        print("  ⚠ Пустой документ")
        return None

    # Состояние парсинга
    author = ""
    title = ""
    description_parts = []
    meta = {}
    ingredient_lines = []
    steps = []          # [{"text": "..."}]
    tip_lines = []
    housewife_lines = []

    # Куда сейчас пишем
    # header → after_table (описание) → ingredients → steps → tip → housewife
    section = "header"
    current_step_text = []
    header_paragraphs_seen = 0  # счётчик НЕ-SEO параграфов в шапке
    saw_table = False

    for kind, content in items:
        if kind == "table":
            # Метаданные из таблицы
            meta.update(parse_table_meta(content))
            saw_table = True
            # После таблицы переходим в режим сбора описания
            if section == "header":
                section = "after_table"
            continue

        # kind == "p"
        text = content

        # Проверяем переключения секций по ключевым словам
        if is_section(text, SECTION_INGREDIENTS):
            # Перед переключением, если был накопленный шаг, сохраним
            if current_step_text:
                steps.append({"text": " ".join(current_step_text).strip()})
                current_step_text = []
            section = "ingredients"
            continue
        if is_section(text, SECTION_STEPS):
            section = "steps"
            continue
        if is_section(text, SECTION_TIP):
            if current_step_text:
                steps.append({"text": " ".join(current_step_text).strip()})
                current_step_text = []
            section = "tip"
            continue
        if is_section(text, SECTION_HOUSEWIFE):
            if current_step_text:
                steps.append({"text": " ".join(current_step_text).strip()})
                current_step_text = []
            section = "housewife"
            continue

        if section == "header":
            # SEO-блок и URL пропускаем
            if is_seo_line(text):
                continue

            # Если встретили группу (Для теста:) — переключаемся на ингредиенты
            group_name = is_group_line(text)
            if group_name:
                section = "ingredients"
                ingredient_lines.append({"__group__": group_name})
                continue

            # Если встретили шаг (1. Текст...) — переключаемся на шаги
            inline_match = STEP_INLINE_RE.match(text)
            if inline_match:
                section = "steps"
                current_step_text.append(inline_match.group(2))
                continue

            # Определяем характер параграфа:
            # - is_long: длинный текст или заканчивается точкой → описание
            # - is_author: 2-3 слова, все с заглавной → имя автора
            # - иначе: кандидат на название
            def looks_like_author(s):
                words = s.split()
                if not (2 <= len(words) <= 3):
                    return False
                return all(w[:1].isupper() for w in words if w)

            is_long = len(text) > 80 or text.endswith((".", "!", "?"))
            is_author_candidate = looks_like_author(text) and not is_long

            if is_long:
                description_parts.append(text)
            elif is_author_candidate and not author:
                # Похоже на автора и автор ещё не заполнен → автор
                author = text
            elif not title:
                # Не похоже на автора (или автор уже есть) и название ещё не заполнено → название
                title = text
            elif not author and is_author_candidate:
                # Резерв: вдруг название было раньше, теперь нашлось имя автора
                author = text
            else:
                # И название, и автор есть — лишний короткий параграф уходит в описание
                description_parts.append(text)
        elif section == "after_table":
            # После таблицы — собираем описание ДО блока ИНГРЕДИЕНТЫ.
            if is_seo_line(text):
                continue
            # Группа → начало ингредиентов без явного заголовка
            group_name = is_group_line(text)
            if group_name:
                section = "ingredients"
                ingredient_lines.append({"__group__": group_name})
                continue
            description_parts.append(text)
        elif section == "ingredients":
            # Внутри секции ингредиентов проверяем — не маркер ли шага?
            # Если да, автоматически переключаемся в режим шагов (если в документе
            # нет заголовка "Рецепт приготовления", а шаги идут сразу после ингредиентов).
            inline_match = STEP_INLINE_RE.match(text)
            if (STEP_HEADER_ONLY_RE.match(text) or
                STEP_NUMBER_ONLY_RE.match(text) or
                inline_match):
                section = "steps"
                if inline_match:
                    current_step_text.append(inline_match.group(2))
            else:
                # Может быть склейка: "Сулугуни — 200 г 1. Нагрейте духовку..."
                # Ищем в строке маркер шага "1. " / "1) " ВНУТРИ текста (не в начале).
                # Берём ПЕРВЫЙ найденный, потому что это начало шагов.
                # Требуем пробел перед номером, чтобы не зацепить "200г1." и т.п.
                split_match = re.search(r"\s(\d+)\s*[\.\)]\s*(.+)$", text)
                if split_match and int(split_match.group(1)) <= 3:
                    # Перед маркером — ингредиент, после — начало шагов
                    before = text[:split_match.start()].rstrip()
                    after = split_match.group(2).strip()
                    if before:
                        ingredient_lines.append(before)
                    section = "steps"
                    if after:
                        current_step_text.append(after)
                else:
                    group_name = is_group_line(text)
                    if group_name:
                        ingredient_lines.append({"__group__": group_name})
                    else:
                        ingredient_lines.append(text)
        elif section == "steps":
            # Маркеры начала нового шага:
            #   "Шаг N" (отдельной строкой)
            #   "N." или "N)" (отдельной строкой)
            #   "N. текст шага..." (inline)
            inline_match = STEP_INLINE_RE.match(text)
            if STEP_HEADER_ONLY_RE.match(text) or STEP_NUMBER_ONLY_RE.match(text):
                # Сохраняем предыдущий шаг (если был), начинаем новый
                if current_step_text:
                    steps.append({"text": " ".join(current_step_text).strip()})
                    current_step_text = []
            elif inline_match:
                # Inline-маркер: "1. Возьмите чеснок..." — сохраняем предыдущий шаг,
                # текст после номера идёт в новый шаг
                if current_step_text:
                    steps.append({"text": " ".join(current_step_text).strip()})
                    current_step_text = []
                current_step_text.append(inline_match.group(2))
            else:
                current_step_text.append(text)
        elif section == "tip":
            tip_lines.append(text)
        elif section == "housewife":
            housewife_lines.append(text)

    # Не забыть последний шаг
    if current_step_text:
        steps.append({"text": " ".join(current_step_text).strip()})

    # Парсим строки ингредиентов (могут быть dict-группы и str-ингредиенты)
    ingredients = []
    for line in ingredient_lines:
        if isinstance(line, dict) and "__group__" in line:
            ingredients.append({"group": line["__group__"]})
        else:
            ing = parse_ingredient_line(line)
            if ing and ing["name"]:
                ingredients.append(ing)

    # Извлекаем картинки: из ZIP (external_images) или из самого docx
    if external_images is not None:
        images = external_images
        print(f"  • Картинок из архива: {len(images)}")
    else:
        img_subdir = images_root / docx_path.stem
        images = extract_images_from_docx(docx_path, img_subdir)
        print(f"  • Картинок извлечено: {len(images)}")

    main_photo = str(images[0].relative_to(images_root)) if images else ""
    # Привязываем картинки к шагам по порядку (первая — главное фото)
    for i, step in enumerate(steps):
        img_idx = i + 1  # сдвиг на 1, т.к. первая — главное фото
        if 0 <= img_idx < len(images):
            step["image"] = str(images[img_idx].relative_to(images_root))
        else:
            step["image"] = ""

    # Название рецепта.
    # Логика: приоритет имени файла (без расширения, "_" → пробел),
    # НО если первая строка документа выглядит как осмысленное название
    # (есть жирные звёздочки **...** или явно короткая строка-заголовок), берём её.
    title_from_doc = re.sub(r"\*+", "", title).strip()
    # Имя файла может содержать фамилию автора: "название_Фамилия.docx"
    # Если автор ещё не найден в документе — пробуем взять из имени файла.
    if not author and "_" in docx_path.stem:
        candidate = docx_path.stem.rsplit("_", 1)[-1].strip()
        if candidate and candidate[0].isupper() and candidate.replace(" ", "").isalpha():
            author = candidate
    title_from_filename = re.sub(r"_[А-ЯЁA-Z][а-яёa-z]+$", "", docx_path.stem).replace("_", " ").strip()

    # Признаки "явного названия" в документе:
    # - окружено звёздочками (Markdown-bold)
    # - не пустое и не очень длинное (до 100 символов — это название, а не описание)
    has_explicit_title = (
        "*" in title  # были жирные звёздочки
        or (title_from_doc and len(title_from_doc) <= 100
            and not title_from_doc.startswith("Описание")
            and title_from_doc != author)
    )
    final_title = title_from_doc if has_explicit_title and title_from_doc else title_from_filename

    recipe = {
        "author": author,
        "title": final_title,
        "description": " ".join(description_parts).strip(),
        "servings": meta.get("servings", 4),
        "cooking_time": meta.get("cooking_time", 30),
        "extra_time": meta.get("extra_time", 0),
        "main_photo": main_photo,
        "ingredients": ingredients,
        "steps": steps,
        "tip": " ".join(tip_lines).strip(),
        "housewife_note": " ".join(housewife_lines).strip(),
        "meta_dropdowns": {
            "main_ingredient": meta.get("main_ingredient", ""),
            "dish_type": meta.get("dish_type", ""),
            "difficulty": meta.get("difficulty", ""),
            "occasion": meta.get("occasion", ""),
            "category": meta.get("category", ""),
            "device": meta.get("device", ""),
            "technology": meta.get("technology", ""),
            "cuisine": meta.get("cuisine", ""),
            "diet_type": meta.get("diet_type", ""),
            "vegetarian_type": meta.get("vegetarian_type", ""),
        },
    }

    print(f"  ✓ Название: {recipe['title'] or '(не нашёл)'}")
    print(f"  • Автор: {author}")
    print(f"  • Ингредиентов: {len(ingredients)}")
    print(f"  • Шагов: {len(steps)}")
    if recipe["tip"]:
        print(f"  • Блок 'Кстати': есть")
    if recipe["housewife_note"]:
        print(f"  • Блок 'Хозяйке на заметку': есть")
    if not ingredients:
        print("  ⚠ Ингредиенты не нашлись — проверь, есть ли заголовок «ИНГРЕДИЕНТЫ»")
    if not steps:
        print("  ⚠ Шаги не нашлись — проверь, есть ли заголовок «РЕЦЕПТ ПРИГОТОВЛЕНИЯ»")

    return recipe


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    target = Path(sys.argv[1])
    if not target.exists():
        print(f"Не нашёл: {target}")
        sys.exit(1)

    import shutil
    import zipfile

    IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".gif"}
    images_root = Path("images")

    def clear_work_dir(work_dir: Path):
        if work_dir.exists():
            for f in work_dir.glob("*.docx"):
                f.unlink()
                print(f"  - удалил: {f.name}")
        else:
            work_dir.mkdir(parents=True)
        if images_root.exists():
            shutil.rmtree(images_root)
            print(f"  - очистил папку картинок: {images_root}/")

    # === РЕЖИМ ZIP АРХИВА ===
    if target.is_file() and target.suffix.lower() == ".zip":
        work_dir = Path("мои-рецепты")
        print(f"Режим: ZIP архив «{target.name}»")
        print(f"Очищаю рабочую папку: {work_dir}/")
        clear_work_dir(work_dir)
        images_root.mkdir(exist_ok=True)

        with zipfile.ZipFile(target) as zf:
            entries = [n for n in zf.namelist() if not n.startswith("__MACOSX")]
            docx_entries = [n for n in entries if Path(n).suffix.lower() == ".docx"]
            img_entries  = sorted(
                [n for n in entries if Path(n).suffix.lower() in IMAGE_EXTS],
                key=lambda x: Path(x).name.lower()
            )

            if not docx_entries:
                print("⚠ В архиве не нашёл .docx файл")
                sys.exit(1)
            if len(docx_entries) > 1:
                print(f"⚠ В архиве несколько .docx, беру первый: {Path(docx_entries[0]).name}")

            docx_entry = docx_entries[0]
            docx_name  = Path(docx_entry).name
            docx_stem  = Path(docx_entry).stem
            docx_target = work_dir / docx_name

            print(f"  + docx: {docx_name}")
            with zf.open(docx_entry) as src, open(docx_target, "wb") as dst:
                dst.write(src.read())

            # Картинки → images/{docx_stem}/  (порядок по имени файла)
            img_subdir = images_root / docx_stem
            img_subdir.mkdir(parents=True, exist_ok=True)
            extracted_images = []
            for entry in img_entries:
                img_filename = Path(entry).name
                img_target   = img_subdir / img_filename
                with zf.open(entry) as src, open(img_target, "wb") as dst:
                    dst.write(src.read())
                extracted_images.append(img_target)
                print(f"  + картинка: {img_filename}")

        docx_files   = [docx_target]
        zip_img_map  = {str(docx_target): extracted_images}  # путь → список Path картинок

    # === РЕЖИМ ОДНОГО .DOCX ===
    elif target.is_file() and target.suffix.lower() == ".docx":
        work_dir = Path("мои-рецепты")
        print(f"Режим: один файл «{target.name}»")
        print(f"Очищаю рабочую папку: {work_dir}/")
        clear_work_dir(work_dir)
        images_root.mkdir(exist_ok=True)

        new_path = work_dir / target.name
        shutil.copy2(target, new_path)
        print(f"  + скопировал: {target.name} → {new_path}")

        docx_files  = [new_path]
        zip_img_map = {}

    # === РЕЖИМ ПАПКИ ===
    else:
        images_root.mkdir(exist_ok=True)
        docx_files  = [target] if target.is_file() else sorted(target.glob("*.docx"))
        zip_img_map = {}

    if not docx_files:
        print(f"Не нашёл .docx файлов в {target}")
        sys.exit(1)

    print(f"\nОбрабатываю файлов: {len(docx_files)}")

    recipes = []
    for path in docx_files:
        try:
            ext_imgs = zip_img_map.get(str(path))  # None если не ZIP-режим
            r = convert_docx(path, images_root, external_images=ext_imgs)
            if r:
                recipes.append(r)
        except Exception as e:
            print(f"  ✗ Ошибка на {path.name}: {e}")
            import traceback
            traceback.print_exc()

    output_file = Path("recipes.json")
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(recipes, f, ensure_ascii=False, indent=2)

    print(f"\n✓ Готово! Записано {len(recipes)} рецептов в {output_file}")
    print(f"✓ Картинки сохранены в папке: {images_root}/")

    # Удаляем исходный ZIP после успешной обработки
    if target.is_file() and target.suffix.lower() == ".zip" and recipes:
        target.unlink()
        print(f"✓ Удалил исходный архив: {target}")

    print("\nЧто проверить руками в recipes.json:")
    print("  1. Правильно ли разобрались названия и ингредиенты")
    print("  2. Правильно ли привязались картинки к шагам")


if __name__ == "__main__":
    main()
