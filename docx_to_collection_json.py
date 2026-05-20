#!/usr/bin/env python3
"""
Конвертер подборок рецептов из Word (.docx) в JSON.

Формат docx:
    Автор
    Заголовок подборки
    Вступительный текст
    Название рецепта 1
    https://www.gastronom.ru/recipe/.../slug
    Описание рецепта 1
    (время, ингредиенты, шаги — игнорируются)
    Название рецепта 2
    ...

Запуск:
    python docx_to_collection_json.py подборка.docx
    python docx_to_collection_json.py папка/с/подборками/
"""

import re
import json
import sys
import base64
import html as html_lib
from pathlib import Path
from urllib.request import urlopen, Request

if sys.platform == "win32":
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
except ImportError:
    print("Нужна библиотека: pip install python-docx")
    sys.exit(1)

RECIPE_URL_RE = re.compile(r'https?://(?:www\.)?gastronom\.ru/recipe/\d+/[\w-]+')

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "ru-RU,ru;q=0.9",
}

SKIP_PREFIXES = (
    "для приготовления",
    "полезный совет",
    "дополнительное время",
    "для глазури",
    "для начинки",
    "для украшения",
    "для соуса",
    "для подачи",
)

STEP_RE = re.compile(r"^\d+\.\s")
TIME_RE  = re.compile(r"^\d+\s*(минут|час|мин)", re.IGNORECASE)


def is_skip(text: str) -> bool:
    t = text.lower().strip()
    if RECIPE_URL_RE.match(t):
        return False  # URL не пропускаем здесь
    if STEP_RE.match(text) or TIME_RE.match(text):
        return True
    for p in SKIP_PREFIXES:
        if t.startswith(p):
            return True
    # Строки-заголовки вида "Для теста:" / "Глазурь:"
    if re.match(r'^[а-яёА-ЯЁ][а-яёА-ЯЁ\s]*:\s*$', text):
        return True
    return False


def is_long(text: str) -> bool:
    return len(text) > 80 or text.endswith((".", "!", "?", "…", "–", "—"))


def get_image_uuid(recipe_url: str) -> str:
    """
    Получает UUID главного изображения рецепта.
    Стратегия: берём og:image, последний сегмент URL — это base64-путь к файлу в CMS.
    Декодируем и вытаскиваем имя файла без расширения (UUID).
    """
    try:
        req = Request(recipe_url, headers=HEADERS)
        html = urlopen(req, timeout=15).read().decode("utf-8", errors="replace")

        m = re.search(r'og:image["\s][^>]*content="([^"]+)"', html)
        if not m:
            m = re.search(r'content="([^"]+)"[^>]*og:image', html)
        if not m:
            return ""

        img_url = m.group(1)
        last_seg = img_url.split("/")[-1]          # base64.webp
        b64_part = last_seg.rsplit(".", 1)[0]       # убираем расширение

        padding = "=" * (-len(b64_part) % 4)
        decoded = base64.b64decode(b64_part + padding).decode("utf-8")
        # decoded = "/cms/all-images/UUID.jpg"
        fname = decoded.split("/")[-1]              # UUID.jpg
        return fname.rsplit(".", 1)[0]              # UUID

    except Exception as e:
        print(f"  ⚠ Не удалось получить UUID картинки для {recipe_url}: {e}")
        return ""


def _postprocess_html(html: str, recipe_names: set | None = None) -> str:
    """
    Пост-обработка HTML от mammoth:
    1. Удаляет строки с URL gastronom.ru (h1-h6 и p)
    2. Названия рецептов → <h2>
    3. Подсекции ингредиентов («Для маринада:» и т.п.) → <p><em>...</em></p>
    4. Прочие заголовки mammoth → <p><strong>...</strong></p>
    5. Ингредиенты (после «Для приготовления») → <ul><li>
    6. Нумерованные шаги (1. / 1) ) → <ol><li>
    7. <ol>/<ul> от mammoth (Word-списки) — сохраняются целиком без разбора
    """
    if recipe_names is None:
        recipe_names = set()
    _names_norm = {n.strip().lower() for n in recipe_names}

    def is_recipe_name(plain_text: str) -> bool:
        return plain_text.strip().lower() in _names_norm

    GASTRO_URL = re.compile(r'https?://(?:www\.)?gastronom\.ru/recipe/\d+/[\w-]+')
    STEP_START  = re.compile(r'^\d+\.\s')
    SECTION_HDR = re.compile(r'^.{1,60}:\s*$')
    ALL_STRONG  = re.compile(r'^<strong>(.+?)</strong>$', re.DOTALL)

    # Вытаскиваем <p>/<h> из внутренностей <ol>/<ul> (mammoth иногда туда помещает)
    def fix_list_nesting(h: str) -> str:
        def repl(m):
            open_tag, inner, close_tag = m.group(1), m.group(2), m.group(3)
            extra = re.findall(r'<(?:p|h[1-6])\b[^>]*>.*?</(?:p|h[1-6])>', inner, re.DOTALL)
            clean = re.sub(r'<(?:p|h[1-6])\b[^>]*>.*?</(?:p|h[1-6])>', '', inner, flags=re.DOTALL)
            return open_tag + clean.strip() + close_tag + ''.join(extra)
        return re.sub(r'(<(?:ol|ul)[^>]*>)(.*?)(</(?:ol|ul)>)', repl, h, flags=re.DOTALL)

    html = fix_list_nesting(html)

    # НЕ разбиваем на <li>, чтобы сохранить <ol>/<ul> от mammoth целыми
    blocks = re.split(r'(?=<(?:h[1-6]|p|ol|ul)\b)', html)

    result = []
    in_ingredient_section = False
    in_ul = False
    step_counter = [0]  # счётчик шагов для mammoth-ol

    def close_ul():
        nonlocal in_ul
        if in_ul:
            result.append('</ul>')
            in_ul = False

    for block in blocks:
        if not block.strip():
            continue

        # ── <ol> от mammoth → нумерованные <p>N. текст</p> ──────────
        if re.match(r'<ol\b', block):
            close_ul()
            in_ingredient_section = False
            li_items = re.findall(r'<li[^>]*>(.*?)</li>', block, re.DOTALL)
            for i, item in enumerate(li_items, 1):
                result.append(f'<p>{i}. {item.strip()}</p>')
            step_counter[0] = len(li_items)
            continue

        # ── <ul> от mammoth — сохраняем как есть ─────────────────────
        if re.match(r'<ul\b', block):
            close_ul()
            in_ingredient_section = False
            result.append(block)
            continue

        # ── Заголовок ─────────────────────────────────────────────────
        if re.match(r'<h[1-6]\b', block):
            close_ul()
            inner = re.sub(r'<[^>]+>', '', block).strip()

            if GASTRO_URL.fullmatch(inner):
                continue

            if is_recipe_name(inner):
                in_ingredient_section = False
                step_counter[0] = 0
                block = re.sub(r'^<h[1-6](\b[^>]*)>', r'<h2\1>', block)
                block = re.sub(r'</h[1-6]>$', r'</h2>', block)
                result.append(block)
            elif in_ingredient_section and SECTION_HDR.match(inner):
                result.append(f'<p><em>{inner}</em></p>')
                in_ingredient_section = True
            else:
                in_ingredient_section = bool(
                    re.search(r'для приготовления', inner, re.IGNORECASE)
                )
                result.append(f'<p><strong>{inner}</strong></p>')
            continue

        # ── Параграф ──────────────────────────────────────────────────
        if re.match(r'<p\b', block):
            m = re.match(r'<p[^>]*>(.*?)</p>\s*$', block, re.DOTALL)
            inner = m.group(1).strip() if m else block
            plain = re.sub(r'<[^>]+>', '', inner).strip()

            if GASTRO_URL.fullmatch(plain):
                continue

            # <p><strong>текст</strong></p>
            strong_m = ALL_STRONG.match(inner)
            if strong_m and len(plain) < 120 and not GASTRO_URL.search(plain):
                if is_recipe_name(plain):
                    close_ul()
                    in_ingredient_section = False
                    step_counter[0] = 0
                    result.append(f'<h2>{strong_m.group(1)}</h2>')
                elif in_ingredient_section and SECTION_HDR.match(plain):
                    close_ul()
                    in_ingredient_section = True
                    result.append(f'<p><em>{plain}</em></p>')
                else:
                    in_ingredient_section = bool(
                        re.search(r'для приготовления', plain, re.IGNORECASE)
                    )
                    result.append(block)
                continue

            # Нумерованный шаг (вручную "1. текст") → <p>N. текст</p>
            if STEP_START.match(plain):
                in_ingredient_section = False
                close_ul()
                result.append(f'<p>{inner}</p>')
                continue

            # Внутри секции ингредиентов
            if in_ingredient_section:
                if SECTION_HDR.match(plain):
                    close_ul()
                    result.append(f'<p><em>{plain}</em></p>')
                else:
                    if not in_ul:
                        result.append('<ul>')
                        in_ul = True
                    result.append(f'<li>{inner}</li>')
                continue

            result.append(block)
            continue

        result.append(block)

    close_ul()
    out = ''.join(result)
    # W-кнопка сайта понимает <b>/<i>, но игнорирует <strong>/<em>
    out = out.replace('<strong>', '<b>').replace('</strong>', '</b>')
    out = out.replace('<em>', '<i>').replace('</em>', '</i>')
    return out


def _mammoth_section_html(docx_path: Path, start_para_index: int,
                          recipe_names: set | None = None) -> str:
    """
    Конвертирует docx в HTML через mammoth, затем обрезает начало до нужного параграфа.
    Избегает копирования параграфов между документами (проблема с rId).
    """
    try:
        import mammoth
    except ImportError:
        return ""

    src = Document(docx_path)

    # Текст первого нужного параграфа — для поиска в HTML
    first_text = ""
    if start_para_index < len(src.paragraphs):
        first_text = src.paragraphs[start_para_index].text.strip()

    # Конвертируем весь docx целиком
    with open(docx_path, "rb") as f:
        full_html = mammoth.convert_to_html(f).value

    # Находим начало нужной секции по тексту первого параграфа
    if first_text:
        needle = re.escape(first_text[:50])
        m = re.search(needle, full_html)
        if m:
            # Ищем открывающий блочный тег (<p>, <h2> и т.п.) перед текстом,
            # шагая назад мимо инлайн-тегов (<strong>, <em>, <a>)
            pos = m.start()
            for _ in range(10):
                tag_start = full_html.rfind('<', 0, pos)
                if tag_start < 0:
                    break
                snippet = full_html[tag_start:tag_start + 15]
                if re.match(r'<(?:strong|em|b\b|i\b|a\b|span)', snippet, re.I):
                    pos = tag_start  # шагаем ещё левее
                else:
                    full_html = full_html[tag_start:]
                    break

    return _postprocess_html(full_html, recipe_names=recipe_names)


def convert_collection(docx_path: Path) -> dict:
    doc = Document(docx_path)
    all_paras = [p for p in doc.paragraphs if p.text.strip()]
    paras = [p.text.strip() for p in all_paras]

    # Позиции всех URL рецептов
    url_positions = [i for i, p in enumerate(paras) if RECIPE_URL_RE.match(p)]
    if not url_positions:
        raise ValueError("В документе не найдено ссылок на рецепты gastronom.ru")

    # Индекс первого рецепта в полном списке doc.paragraphs (не only non-empty)
    all_doc_paras = doc.paragraphs
    non_empty_to_full = [i for i, p in enumerate(all_doc_paras) if p.text.strip()]
    first_recipe_full_idx = non_empty_to_full[url_positions[0] - 1]

    # Собираем имена рецептов заранее — нужны для выбора h2 vs h4
    recipe_names: set = set()
    for url_pos in url_positions:
        for i in range(url_pos - 1, -1, -1):
            candidate = paras[i]
            if candidate and not RECIPE_URL_RE.match(candidate) and not is_skip(candidate):
                recipe_names.add(candidate)
                break

    # HTML с полным форматированием через mammoth
    body_html = _mammoth_section_html(docx_path, first_recipe_full_idx, recipe_names)
    if not body_html:
        # Фолбек если mammoth не установлен
        print("  ⚠ mammoth не установлен, HTML будет без форматирования (pip install mammoth)")
        body_html = "".join(
            f"<{'h2' if all(r.bold for r in p.runs if r.text.strip()) and len(p.text) < 120 else 'p'}>"
            f"{html_lib.escape(p.text.strip())}"
            f"</{'h2' if all(r.bold for r in p.runs if r.text.strip()) and len(p.text) < 120 else 'p'}>"
            for p in all_paras[url_positions[0] - 1:]
            if p.text.strip()
        )

    # Текст (plain) — для информации
    body_text = "\n\n".join(paras[url_positions[0] - 1:])

    # ── Шапка (до первого URL) ───────────────────────────────────────────────
    header = paras[:url_positions[0]]
    author = ""
    title  = ""
    intro_parts = []

    for p in header:
        if is_skip(p):
            continue
        words = p.split()
        is_name_like = (
            2 <= len(words) <= 3
            and all(w[:1].isupper() for w in words if w)
            and not p.endswith(".")
        )
        if not author and is_name_like and not title:
            # Первая строка — может быть автором или заголовком
            # Если заголовок уже есть — это автор; иначе смотрим следующую логику
            author = p
        elif not title:
            title = p
        elif is_long(p) or len(p) > 40:
            intro_parts.append(p)

    # Если автор и заголовок перепутались — исправим эвристикой
    if author and not title:
        title, author = author, ""

    intro = " ".join(intro_parts).strip()

    # ── Рецепты ──────────────────────────────────────────────────────────────
    recipes = []
    for idx, url_pos in enumerate(url_positions):
        url = paras[url_pos]

        # Название: ближайшая непустая строка ПЕРЕД URL (не URL, не skip)
        name = ""
        for i in range(url_pos - 1, -1, -1):
            candidate = paras[i]
            if candidate and not RECIPE_URL_RE.match(candidate) and not is_skip(candidate):
                name = candidate
                break

        # Описание: первый длинный абзац ПОСЛЕ URL (до следующего URL)
        description = ""
        end = url_positions[idx + 1] if idx + 1 < len(url_positions) else len(paras)
        for i in range(url_pos + 1, end):
            p = paras[i]
            if is_skip(p) or RECIPE_URL_RE.match(p):
                continue
            if is_long(p) or len(p) > 50:
                description = p
                break

        recipes.append({
            "name": name,
            "url": url,
            "description": description,
            "image_uuid": "",   # заполняется отдельно через get_image_uuid()
        })

    return {
        "author": author,
        "title": title,
        "intro": intro,
        "body_html": body_html,
        "body_text": body_text,
        "recipes": recipes,
    }


def fetch_uuids(collection: dict) -> None:
    """Обходит все рецепты и подгружает UUID главного изображения."""
    recipes = collection["recipes"]
    for i, r in enumerate(recipes):
        print(f"  [{i+1}/{len(recipes)}] {r['name'][:50]} → получаю UUID картинки...")
        r["image_uuid"] = get_image_uuid(r["url"])
        if r["image_uuid"]:
            print(f"         UUID: {r['image_uuid']}")
        else:
            print(f"         ⚠ UUID не получен — заполни вручную")


def main():
    if len(sys.argv) < 2:
        print("Использование:")
        print("  python docx_to_collection_json.py подборка.docx")
        print("  python docx_to_collection_json.py папка/")
        sys.exit(1)

    target = Path(sys.argv[1])
    files = []
    if target.is_dir():
        files = sorted(target.glob("*.docx"))
        if not files:
            print(f"Нет .docx файлов в {target}")
            sys.exit(1)
    elif target.suffix.lower() == ".docx":
        files = [target]
    else:
        print(f"Ожидается .docx файл или папка, получено: {target}")
        sys.exit(1)

    collections = []
    for docx_path in files:
        print(f"\nОбрабатываю: {docx_path.name}")
        try:
            col = convert_collection(docx_path)
            print(f"  Заголовок : {col['title']}")
            print(f"  Автор     : {col['author']}")
            print(f"  Рецептов  : {len(col['recipes'])}")
            for r in col["recipes"]:
                print(f"    - {r['name']}")

            print("  Подгружаю UUID картинок...")
            fetch_uuids(col)
            collections.append(col)
        except Exception as e:
            print(f"  ✗ Ошибка: {e}")

    if not collections:
        print("Ничего не сконвертировано.")
        sys.exit(1)

    out = Path("collections.json")
    with open(out, "w", encoding="utf-8") as f:
        json.dump(collections, f, ensure_ascii=False, indent=2)

    total = sum(len(c["recipes"]) for c in collections)
    print(f"\nГотово! Подборок: {len(collections)}, рецептов: {total}. Сохранено в {out}")


if __name__ == "__main__":
    main()
